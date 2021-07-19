from datetime import date
from openpyxl import load_workbook
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib import colors
from docx import Document

file_name_excel = 'Excel_And_Pdf/Customers Data Base.xlsx'


# TODO: Получить пример PDF'a, сконифгурировать шаблон и его пересохранять.


#def save_as_pdf(user_data: dict, index: str) -> str:
#    file_name_pdf = index + "_" + user_data['Имя клиента'] + '.pdf'
#    text_lines = [
#        'Name: ' + user_data['Имя клиента'],
#        'Phone Number: ' + user_data['Номер телефона'],
#        'Issue: ' + user_data['Поломка']
#    ]
#    pdf = canvas.Canvas('Excel_And_Pdf/PDF/' + file_name_pdf)
#    return file_name_pdf


def assignedDataToExcel(user_data, repair_number: str, row: int, done: bool) -> None:
    """
    Задаём в таблицу In_Progress значения по строке + 1 - если done == False
    Задаём в таблицу Done значения из In_Progress - если done == True
    """
    wb = load_workbook(file_name_excel)
    In_Progress_Table = wb['In_Progress']
    Repair_Done = wb['Done']
    if not done:
        In_Progress_Table["A" + str(row + 1)] = repair_number
        In_Progress_Table["B" + str(row + 1)] = user_data['Имя клиента']
        In_Progress_Table["C" + str(row + 1)] = user_data['Номер телефона']
        In_Progress_Table["D" + str(row + 1)] = user_data['Поломка']
    else:
        Repair_Done["A" + str(row + 1)] = repair_number
        Repair_Done["B" + str(row + 1)] = user_data[0].value
        Repair_Done["C" + str(row + 1)] = user_data[1].value
        Repair_Done["D" + str(row + 1)] = user_data[2].value
    wb.save(file_name_excel)


def store_file(user_data) -> str:
    today = date.today()
    wb = load_workbook(file_name_excel)
    sheet = wb.active
    last_row = sheet.max_row
    last_row_value = sheet.cell(column=1, row=last_row).value
    if today.strftime('%d%m%y') == last_row_value[:6]:
        counter = int(last_row_value[6:]) + 1  # Если дата всё таже, то присваиваем последний номер + 1
    else:
        counter = 1  # Если дата изменилась, то обнуляем счётчик
    repair_num = today.strftime('%d%m%y') + str(counter).zfill(2)  # Создаём номер ремонта
    assignedDataToExcel(user_data=user_data, repair_number=repair_num, row=last_row, done=False)
    return create_word(repair_num, user_data['Имя клиента'], user_data['Поломка'])


# TODO: Сделать файл в Google Drive, чтобы в него можно было писать и бот подхватывал это.


def save_data_to_another_table(repair_number: str) -> bool:
    wb = load_workbook(file_name_excel)
    In_Progress = wb['In_Progress']
    Done_Repair = wb['Done']
    for row_number in range(1, len(In_Progress['A'])):
        if In_Progress["A"][row_number].value == repair_number:  # Проверяем, есть ли такой ремонт в таблице
            assignedDataToExcel(user_data=In_Progress[row_number + 1][1:], repair_number=repair_number,
                                row=Done_Repair.max_row, done=True)
            In_Progress.delete_rows(row_number + 1, 1)  # Удаляем строку из In_Progress
            wb.save(file_name_excel)
            break
    return True


def create_word(repair_num: str, customer_name: str, customer_problem: str) -> str:
    test = Document('Template_Customer.docx')
    file_name_doc = repair_num + "_" + customer_name + '.docx'
    dict_customer = {
        1: f'Remondi vastuvõtmise kviitung nr {repair_num}, kuupäev {repair_num[:-2]}',
        2: f'{customer_name}\t\t\t(Eesnimi, perekonnanimi, ID)',
        6: f'{customer_problem}'
    }
    for table_num in [0, 2]:
        for row in [1, 2, 6]:
            test.tables[table_num].cell(row, 2).paragraphs[0].add_run(dict_customer[row]).bold = True
    test.save('Excel_And_Pdf/PDF/' + file_name_doc)
    return file_name_doc
